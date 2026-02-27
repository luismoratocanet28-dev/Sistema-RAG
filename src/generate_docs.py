from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def create_work_report():
    doc = Document()
    
    # Título Principal
    title = doc.add_heading('Informe de Desarrollo: Sistema RAG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fecha y Autor
    p = doc.add_paragraph()
    p.add_run(f'Fecha: {datetime.datetime.now().strftime("%d/%m/%Y")}\n').bold = True
    p.add_run('Desarrollado por: Luis Morato Canet (asistido por Antigravity AI)').italic = True
    
    # 1. Proceso de Creación
    doc.add_heading('1. Proceso de Creación', level=1)
    doc.add_paragraph(
        "El proceso comenzó con la interacción inicial con el asistente de IA Antigravity. "
        "A continuación se detallan los pasos realizados:"
    )
    
    steps = [
        "Acceso a la plataforma Antigravity para solicitar la creación del sistema.",
        "Definición de requerimientos: sistema de recuperación y generación, documentación técnica y subida a GitHub.",
        "Investigación de tecnologías adecuadas para un sistema RAG moderno y eficiente.",
        "Configuración del entorno local en el PC del usuario (Python 3.14).",
        "Implementación del código fuente utilizando LangChain para la orquestación y FAISS para el almacenamiento vectorial.",
        "Generación automática del presente documento para registrar el flujo de trabajo.",
        "Preparación para el despliegue en el repositorio personal de GitHub."
    ]
    
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')
        
    # 2. Posibilidades y Alternativas Consideradas
    doc.add_heading('2. Posibilidades y Alternativas', level=1)
    doc.add_paragraph(
        "Durante la fase de diseño, se evaluaron diversas arquitecturas para el sistema RAG:"
    )
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Opción'
    hdr_cells[1].text = 'Ventajas'
    hdr_cells[2].text = 'Desventajas'
    
    opts = [
        ("SaaS Total (Azure/OpenAI + Pinecone)", "Escalabilidad total, sin mantenimiento hardware.", "Coste recurrente elevado, privacidad cedida a terceros."),
        ("Local Open Source (Ollama + FAISS)", "Privacidad absoluta, coste cero.", "Requiere hardware potente (GPU) y es más complejo de configurar."),
        ("Solución Híbrida (LangChain + Local Embeddings + GPT)", "Equilibrio entre coste y rendimiento. Fácil de usar y compatible con Python 3.14.", "Requiere una clave de API para la generación final.")
    ]
    
    for opt, advantage, disadvantage in opts:
        row_cells = table.add_row().cells
        row_cells[0].text = opt
        row_cells[1].text = advantage
        row_cells[2].text = disadvantage

    # 3. Solución Escogida
    doc.add_heading('3. Solución Escogida y Justificación', level=1)
    doc.add_paragraph(
        "Se ha seleccionado la Solución Híbrida. Los motivos principales son:"
    )
    justifications = [
        "Eficiencia de Costes: Al usar embeddings locales (HuggingFace), no se paga por la indexación de documentos.",
        "Simplicidad y Compatibilidad: FAISS funciona como una base de datos local ligera y ha demostrado ser más estable en entornos con Python 3.14.",
        "Flexibilidad: El sistema permite cambiar fácilmente de modelo de lenguaje (LLM) modificando una sola línea de código.",
        "Privacidad inicial: Los documentos se procesan localmente antes de cualquier consulta externa."
    ]
    for item in justifications:
        doc.add_paragraph(item, style='List Bullet')

    # Guardar documento
    filename = "Proceso_Creacion_Sistema_RAG.docx"
    doc.save(filename)
    print(f"Documento '{filename}' creado con éxito.")

if __name__ == "__main__":
    create_work_report()
